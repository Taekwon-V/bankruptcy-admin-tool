import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def create_interview_docx():
    doc = Document()
    
    # 1. Page setup - A4, normal margins (20mm / ~0.79 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    
    # Colors
    NAVY = RGBColor(27, 54, 93)       # #1B365D
    DARK_GRAY = RGBColor(40, 40, 40)  # #282828
    BLUE = RGBColor(37, 99, 235)      # #2563EB
    MUTED = RGBColor(100, 116, 139)   # #64748B
    
    # Helper: Add Korean Run with explicit EastAsia font (fixes tofu/box rendering in MS Word)
    def add_ko_run(paragraph, text, font_name='맑은 고딕', size_pt=10.5, bold=False, color=None):
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
            
        # Crucial for MS Word Korean / EastAsia text rendering:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        return run

    # Helper: Set cell background color
    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
    # Helper: Set cell padding
    def set_cell_padding(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Set normal style default fonts
    normal_style = doc.styles['Normal']
    normal_style.font.name = '맑은 고딕'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = DARK_GRAY
    normal_style.paragraph_format.line_spacing = 1.35
    normal_style.paragraph_format.space_after = Pt(4)
    rPr = normal_style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    add_ko_run(title_p, "파산관제인 업무 자동화 시스템 요구사항 인터뷰 체크리스트", font_name='맑은 고딕', size_pt=17, bold=True, color=NAVY)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    add_ko_run(sub_p, "실무자 맞춤형 비즈니스 & 기능 중심 인터뷰 가이드 (총 20문항)", font_name='맑은 고딕', size_pt=10.5, color=MUTED)
    
    # Overview Box
    info_table = doc.add_table(rows=1, cols=1)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Inches(6.57)
    
    cell = info_table.rows[0].cells[0]
    set_cell_shading(cell, "F1F5F9")
    set_cell_padding(cell, top=140, bottom=140, left=180, right=180)
    
    info_p = cell.paragraphs[0]
    info_p.paragraph_format.space_after = Pt(2)
    add_ko_run(info_p, "📌 안내 및 목적\n", size_pt=10, bold=True, color=NAVY)
    add_ko_run(
        info_p,
        "본 체크리스트는 파산관제인(개인·법인) 실무에서 겪는 가장 번거롭고 시간이 많이 소요되는 병목 업무를 파악하고, "
        "이를 100% 로컬(PC 독립 구동) 프로그램으로 자동화하기 위한 요구사항 인터뷰 양식입니다. "
        "IT 전문 지식 없이도 평소 실무 관행과 애로사항을 편하게 작성하시거나 답변하실 수 있도록 구성되었습니다.",
        size_pt=9.5, color=DARK_GRAY
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Categories & Questions
    categories = [
        {
            "cat_num": "1",
            "title": "사건 현황 및 주요 병목 업무 (3문항)",
            "desc": "현재 전반적인 업무량과 가장 시간 소모가 큰 단계를 파악합니다.",
            "questions": [
                ("1", "한 달 평균 진행하시는 파산 사건(개인파산 / 법인파산)은 대략 몇 건 정도 되시나요?"),
                ("2", "사건 1건을 수임하여 서류를 처음 열어보고 최종 보고서를 법원에 제출하기까지, 시간과 손이 가장 많이 가는 단계는 구체적으로 어디인가요?"),
                ("3", "현재 관제인(변호사)님과 사무원(실무 직원) 사이에서 업무 분담(누가 어디까지 초안을 잡고, 누가 최종 확인하는지)은 어떻게 이루어지나요?")
            ]
        },
        {
            "cat_num": "2",
            "title": "통장 및 금융거래내역 분석 (핵심 기능) (4문항)",
            "desc": "수년 치 거래내역 검토 시 필요한 분석 룰과 수작업 방식을 파악합니다.",
            "questions": [
                ("4", "채무자가 제출하는 3~5년 치 은행 계좌 거래내역은 주로 어떤 형태(은행 엑셀 파일, 스캔 PDF, 종이 출력물)로 받으시나요?"),
                ("5", "거래내역을 검토하실 때 '이 금액 이상은 무조건 소명을 요구한다'는 기준 금액(예: 건당 300만 원, 500만 원, 1,000만 원 등)이 정해져 있으신가요?"),
                ("6", "고액 거래 외에 반드시 찾아내야 하는 특정 거래처나 키워드(예: 가족/친인척 이름, 가상화폐 거래소, 도박/토토, 보험 해약환급금, 사채/대부업체 등)는 어떤 것들인가요?"),
                ("7", "파산 직전에 특정 채권자나 지인에게만 돈을 몰아 갚은 내역(편파변제)이나 가족에게 보낸 돈을 찾을 때, 현재는 어떤 방식으로 대조하고 계신가요?")
            ]
        },
        {
            "cat_num": "3",
            "title": "필수 제출 서류 및 보정명령 관리 (2문항)",
            "desc": "서류 누락 체크 및 보정 요청서 작성 실무를 파악합니다.",
            "questions": [
                ("8", "채무자가 제출해야 하는 필수 서류(14~20여 종) 중 누락되거나 유효기간이 지나서 다시 제출을 요구(보정 요청)하는 일이 얼마나 자주 발생하나요?"),
                ("9", "보정 요청서나 소명 안내문을 발송할 때, 사무실에서 주로 사용하는 표준 양식이나 자주 쓰는 문구가 따로 있으신가요?")
            ]
        },
        {
            "cat_num": "4",
            "title": "재산 가액 및 환가/면제금액 계산 (2문항)",
            "desc": "법적 면제재산 및 파산재단 환가 대상액 산출 기준을 파악합니다.",
            "questions": [
                ("10", "전세·월세 보증금에서 법적 면제액(소액임차보증금 최우선변제액)을 공제할 때, 지역별 기준표를 직접 찾아보시나요, 아니면 자동으로 계산되면 편하실까요?"),
                ("11", "보험 해약환급금, 예적금, 중고차, 부동산 등에서 '파산재단에 편입시켜 빚 갚는 데 써야 할 금액'을 계산할 때 실무상 적용하는 표준 기준이 있으신가요?")
            ]
        },
        {
            "cat_num": "5",
            "title": "법인파산 특화 검토 (법인 사건 진행 시) (2문항)",
            "desc": "법인 회생/파산 시 재무제표 및 대표자 거래 분석 방식을 확인합니다.",
            "questions": [
                ("12", "법인파산 사건의 경우, 회사의 재무제표(대차대조표 등)에서 대표이사가 가져간 돈(가지급금), 빌려준 돈(가수금), 친인척 거래를 추적하는 작업이 얼마나 자주 필요한가요?"),
                ("13", "회사의 남은 자산(기계, 재고 상품, 외상 매출금 등)을 파악하고 환가 계획을 세울 때 어떤 점이 제일 번거로우신가요?")
            ]
        },
        {
            "cat_num": "6",
            "title": "법원 제출 보고서 및 배당표 작성 (3문항)",
            "desc": "법원 제출 최종 산출물 포맷 및 작성 방식을 파악합니다.",
            "questions": [
                ("14", "법원에 제출하는 '파산관제인 조사보고서'는 주로 한글(HWP/HWPX) 파일로 작성하시나요, 아니면 전자소송 사이트에 직접 입력하시나요?"),
                ("15", "현재 사용하시는 법원 제출용 보고서 빈 양식(한글 파일 샘플) 1부를 제공해 주실 수 있나요?"),
                ("16", "채권자들에게 돈을 나누어 주는 '배당표'나 '채권조사표'를 엑셀이나 한글 표로 직접 손으로 계산하여 작성하고 계신가요?")
            ]
        },
        {
            "cat_num": "7",
            "title": "일정 관리 및 시스템 사용 환경 (4문항)",
            "desc": "기일 관리 방식, 사무실 PC 환경 및 킬러 기능을 확인합니다.",
            "questions": [
                ("17", "채권자집회 기일, 의견서 제출 기한, 보정 기한 등 사건 일정을 현재는 어떻게 관리(달력, 다이어리, 엑셀 등)하고 계신가요?"),
                ("18", "사건별로 채무자 정보, 제출 서류, 통장 분석 결과를 한곳에 사건별 카드/폴더처럼 모아서 검색하고 관리하는 화면이 필요하신가요?"),
                ("19", "사무실 내 여러 컴퓨터에서 데이터를 함께 봐야 하나요, 아니면 관제인/담당자 PC 1~2대에서 독립적으로 사용하면 충분한가요?"),
                ("20", "★ [핵심] '이 기능 딱 하나만 자동으로 1초 만에 해결해 주면 당장이라도 쓰겠다'고 생각하시는 가장 간절한 기능은 무엇인가요?")
            ]
        }
    ]

    for cat in categories:
        # Category Header
        h_p = doc.add_paragraph()
        h_p.paragraph_format.space_before = Pt(12)
        h_p.paragraph_format.space_after = Pt(2)
        h_p.paragraph_format.keep_with_next = True
        
        add_ko_run(h_p, f"■ {cat['cat_num']}. {cat['title']}", size_pt=11.5, bold=True, color=NAVY)
        
        # Category Desc
        d_p = doc.add_paragraph()
        d_p.paragraph_format.space_before = Pt(0)
        d_p.paragraph_format.space_after = Pt(5)
        d_p.paragraph_format.keep_with_next = True
        add_ko_run(d_p, f"   ※ {cat['desc']}", size_pt=9.5, color=MUTED)
        
        # Questions Table / Box
        for q_num, q_text in cat['questions']:
            q_table = doc.add_table(rows=2, cols=1)
            q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            q_table.autofit = False
            q_table.columns[0].width = Inches(6.57)
            
            # Row 0: Question Text
            cell_q = q_table.rows[0].cells[0]
            set_cell_shading(cell_q, "F8FAFC")
            set_cell_padding(cell_q, top=70, bottom=50, left=120, right=120)
            
            qp = cell_q.paragraphs[0]
            qp.paragraph_format.space_before = Pt(0)
            qp.paragraph_format.space_after = Pt(0)
            
            q_color = BLUE if "20" not in q_num else RGBColor(220, 38, 38)
            add_ko_run(qp, f"Q{q_num}. ", size_pt=10, bold=True, color=q_color)
            add_ko_run(qp, q_text, size_pt=10, bold=True, color=DARK_GRAY)
            
            # Row 1: Answer Memo Space
            cell_a = q_table.rows[1].cells[0]
            set_cell_shading(cell_a, "FFFFFF")
            set_cell_padding(cell_a, top=50, bottom=80, left=120, right=120)
            
            ap = cell_a.paragraphs[0]
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after = Pt(0)
            add_ko_run(ap, "답변 / 메모:\n\n", size_pt=9.5, color=RGBColor(160, 174, 192))
            
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Footer note
    doc.add_paragraph().paragraph_format.space_before = Pt(10)
    footer_box = doc.add_table(rows=1, cols=1)
    footer_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_box.columns[0].width = Inches(6.57)
    f_cell = footer_box.rows[0].cells[0]
    set_cell_shading(f_cell, "EFF6FF")
    set_cell_padding(f_cell, top=90, bottom=90, left=140, right=140)
    
    fp = f_cell.paragraphs[0]
    add_ko_run(fp, "💡 미팅 팁: ", size_pt=9.5, bold=True, color=NAVY)
    add_ko_run(
        fp,
        "고객사에서 실제 사용하는 [법원 보고서 한글(HWP) 서식 빈 파일]과 [개인정보를 마스킹한 통장 엑셀 파일 샘플]을 1부씩 전달받으시면 자동화 엔진 설계가 매우 신속하게 진행됩니다.",
        size_pt=9.5, color=DARK_GRAY
    )

    output_path = "파산관제인_업무자동화_인터뷰_체크리스트_v2.docx"
    doc.save(output_path)
    print(f"File created successfully: {output_path}")

if __name__ == "__main__":
    create_interview_docx()
